"""
RecoverIQ - Ingestion File Extractors
Dedicated, robust extractors for CSV, XLSX, JSON, TXT, PDF, and DOCX.
Extracts raw record dictionaries strictly from file contents.
"""

import io
import csv
import json
import re
from typing import List, Dict, Any, Tuple


# ==========================================
# 1. CSV Extractor
# ==========================================
def extract_csv(content: str) -> List[Dict[str, Any]]:
    """Extracts raw payment records from CSV content."""
    if not content or not content.strip():
        return []

    records = []
    # Auto-detect dialect or default to comma/tab
    sample = content[:2048]
    delimiter = ","
    if "\t" in sample and "," not in sample:
        delimiter = "\t"
    elif ";" in sample and "," not in sample:
        delimiter = ";"

    reader = csv.DictReader(io.StringIO(content), delimiter=delimiter)
    for row in reader:
        if not row or not any(str(v).strip() for v in row.values() if v is not None):
            continue
        # Clean keys
        clean_row = {str(k).strip().lower(): v for k, v in row.items() if k is not None}
        records.append(clean_row)

    return records


# ==========================================
# 2. XLSX Extractor (openpyxl)
# ==========================================
def extract_xlsx(file_bytes: bytes) -> List[Dict[str, Any]]:
    """Extracts raw payment records from Excel workbook."""
    if not file_bytes:
        return []

    try:
        import openpyxl
    except ImportError:
        return []

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    records = []

    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows or len(rows) < 2:
            continue

        # Find header row (first row with non-empty cells)
        header_idx = -1
        headers = []
        for idx, r in enumerate(rows):
            non_empty = [str(c).strip().lower() for c in r if c is not None and str(c).strip()]
            # Look for payment / id / amount keywords
            if any("pay" in c or "id" in c or "amount" in c or "txn" in c for c in non_empty):
                header_idx = idx
                headers = [str(c).strip().lower() if c is not None else f"col_{i}" for i, c in enumerate(r)]
                break

        if header_idx == -1:
            # Fallback to first row
            header_idx = 0
            headers = [str(c).strip().lower() if c is not None else f"col_{i}" for i, c in enumerate(rows[0])]

        for row in rows[header_idx + 1:]:
            if not row or not any(c is not None and str(c).strip() for c in row):
                continue
            row_dict = {}
            for i, val in enumerate(row):
                if i < len(headers):
                    row_dict[headers[i]] = val
            records.append(row_dict)

    return records


# ==========================================
# 3. JSON Extractor
# ==========================================
def extract_json(content: str) -> List[Dict[str, Any]]:
    """Extracts raw payment records from JSON content."""
    if not content or not content.strip():
        return []

    try:
        data = json.loads(content)
    except Exception:
        return []

    if isinstance(data, list):
        return [item for item in data if isinstance(item, dict)]

    if isinstance(data, dict):
        # Look for nested payment arrays
        for k in ("payments", "transactions", "records", "data", "items", "incidents"):
            if k in data and isinstance(data[k], list):
                return [item for item in data[k] if isinstance(item, dict)]

        # If dict itself represents a single payment
        if any("pay" in k.lower() or "id" in k.lower() for k in data.keys()):
            return [data]

    return []


from .schema import clean_formatting_delimiters

# ==========================================
# 4. TXT Extractor
# ==========================================
def extract_txt(content: str) -> List[Dict[str, Any]]:
    """
    Extracts payment records from TXT files.
    Supports key-value blocks (Payment ID: ...), pipe-delimited, and CSV-like text.
    Strips decorative delimiters (----, ====, ____) from failure reasons and problems.
    A new record begins strictly when a new Payment ID is encountered.
    """
    if not content or not content.strip():
        return []

    lines = [line.strip() for line in content.splitlines() if line.strip()]
    if not lines:
        return []

    # Check if lines look like pipe or comma delimited table
    if all("|" in l for l in lines[:3]):
        # Pipe delimited
        header_cols = [c.strip().lower() for c in lines[0].split("|") if c.strip()]
        records = []
        for l in lines[1:]:
            # Skip separator lines like |---|---|
            if set(l.replace("|", "").strip()) <= {"-", "=", "_", " ", "\t"}:
                continue
            parts = [clean_formatting_delimiters(c) for c in l.split("|") if c.strip() or len(l.split("|")) > 2]
            if l.startswith("|") and parts and parts[0] == "":
                parts.pop(0)
            if l.endswith("|") and parts and parts[-1] == "":
                parts.pop()
            row_dict = {}
            for i, p in enumerate(parts):
                if i < len(header_cols):
                    row_dict[header_cols[i]] = p
            if row_dict:
                records.append(row_dict)
        if records:
            return records

    # Key-Value Block Parser
    records = []
    current_record = {}
    
    pid_start_re = re.compile(r"(?i)^(payment\s*id|transaction\s*id|txn\s*id|payment|id)\s*[:=\-]\s*([a-zA-Z0-9_\-]+)", re.IGNORECASE)
    generic_pid_re = re.compile(r"^(pay_\w+|PAY_\w+|TXN_\w+|txn_\w+)", re.IGNORECASE)

    for line in lines:
        # Ignore pure decorative delimiter / separator lines (e.g. '--------------------', '====================')
        if not clean_formatting_delimiters(line):
            continue

        start_match = pid_start_re.match(line)
        gen_match = generic_pid_re.match(line)

        if start_match:
            if current_record and ("payment_id" in current_record or "id" in current_record):
                records.append(current_record)
                current_record = {}
            current_record["payment_id"] = clean_formatting_delimiters(start_match.group(2))
            continue
        elif gen_match and not any(k in line.lower() for k in ("amount", "status", "problem", "reason")):
            if current_record and ("payment_id" in current_record or "id" in current_record):
                records.append(current_record)
                current_record = {}
            current_record["payment_id"] = clean_formatting_delimiters(gen_match.group(1))
            continue

        # Parse key-value lines for current record
        if ":" in line or "=" in line:
            sep = ":" if ":" in line else "="
            k, v = line.split(sep, 1)
            k_clean = k.strip().lower()
            v_clean = clean_formatting_delimiters(v)
            if "amount" in k_clean:
                current_record["amount"] = v_clean
            elif "problem" in k_clean or "issue" in k_clean or "failure" in k_clean or "error" in k_clean:
                current_record["problem"] = v_clean
            elif "root" in k_clean or "cause" in k_clean:
                current_record["root_cause"] = v_clean
            elif "status" in k_clean:
                current_record["status"] = v_clean
            elif "decision" in k_clean or "recommendation" in k_clean or "action" in k_clean:
                current_record["recommendation"] = v_clean
            else:
                current_record[k_clean] = v_clean
        elif current_record:
            clean_add = clean_formatting_delimiters(line)
            if clean_add:
                if "problem" in current_record:
                    current_record["problem"] = (current_record.get("problem", "") + " " + clean_add).strip()
                elif "root_cause" in current_record:
                    current_record["root_cause"] = (current_record.get("root_cause", "") + " " + clean_add).strip()

    if current_record and ("payment_id" in current_record or "id" in current_record):
        records.append(current_record)

    return records


# ==========================================
# 5. PDF Extractor (pypdf)
# ==========================================
def extract_pdf(file_bytes: bytes, content_str: Optional[str] = None) -> List[Dict[str, Any]]:
    """
    Extracts payment records from PDF bytes or extracted text.
    Identifies payment ID boundaries, handles multi-line table wrapping,
    and combines lines into exact single payment records without duplication.
    """
    full_text = ""
    if file_bytes and file_bytes.startswith(b"%PDF"):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                txt = page.extract_text()
                if txt:
                    full_text += "\n" + txt
        except Exception:
            pass

    if not full_text and content_str:
        full_text = content_str
    elif not full_text and file_bytes:
        try:
            full_text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            pass

    if not full_text or not full_text.strip():
        return []

    # First attempt: structured TXT block extraction on full text
    txt_records = extract_txt(full_text)
    if txt_records and len(txt_records) > 0:
        return txt_records

    # Second attempt: Regex row scanning for payment ID patterns (e.g. pay_101, pay_001)
    lines = [l.strip() for l in full_text.splitlines() if l.strip()]
    records = []
    current_item = None

    pid_pattern = re.compile(r"\b(pay_\w+|PAY_\w+|TXN_\w+|txn_\w+|pay\d+)\b", re.IGNORECASE)
    amount_pattern = re.compile(r"(?:₹|Rs\.?|INR|\$)?\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?|\d+)\b")

    for line in lines:
        # Ignore page header/footer artifacts
        if re.match(r"(?i)^(page\s*\d+|recoveriq|payment\s*recovery\s*report|confidential)", line):
            continue

        match = pid_pattern.search(line)
        if match:
            # Start of a new payment row
            if current_item:
                records.append(current_item)

            pid = match.group(1)
            # Remove PID from line to extract amount and problem
            line_rem = line[:match.start()] + line[match.end():]
            
            # Extract amount from remainder of line
            amt_match = amount_pattern.search(line_rem)
            amt_val = amt_match.group(0) if amt_match else None
            prob_rem = line_rem
            if amt_match:
                prob_rem = line_rem[:amt_match.start()] + line_rem[amt_match.end():]

            prob_text = prob_rem.strip().strip("|,-: ")

            current_item = {
                "payment_id": pid,
                "amount": amt_val,
                "problem": prob_text
            }
        elif current_item:
            # This is a continuation / wrapped line of the previous payment
            amt_match = amount_pattern.search(line)
            if amt_match and not current_item.get("amount"):
                current_item["amount"] = amt_match.group(0)
                line_no_amt = line[:amt_match.start()] + line[amt_match.end():]
                clean_add = line_no_amt.strip().strip("|,-: ")
                if clean_add:
                    current_item["problem"] = (current_item.get("problem", "") + " " + clean_add).strip()
            else:
                clean_line = line.strip().strip("|,-: ")
                if clean_line:
                    current_item["problem"] = (current_item.get("problem", "") + " " + clean_line).strip()

    if current_item:
        records.append(current_item)

    return records


# ==========================================
# 6. DOCX Extractor (python-docx)
# ==========================================
def extract_docx(file_bytes: bytes) -> List[Dict[str, Any]]:
    """Extracts payment records from Word DOCX tables and paragraphs."""
    if not file_bytes:
        return []

    try:
        import docx
    except ImportError:
        return []

    doc = docx.Document(io.BytesIO(file_bytes))
    records = []

    # 1. Inspect tables first
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            row_cells = [cell.text.strip() for cell in row.cells]
            if not any(row_cells):
                continue
            row_dict = {}
            for idx, cell_text in enumerate(row_cells):
                if idx < len(header_cells):
                    row_dict[header_cells[idx]] = cell_text
            if row_dict:
                records.append(row_dict)

    if records:
        return records

    # 2. Fallback: inspect paragraphs as TXT
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return extract_txt(full_text)
